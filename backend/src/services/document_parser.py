"""
Document Parsing Service
Extracts text from various document formats (PDF, DOCX, HTML, TXT, MD)
"""
from typing import List, Dict, Any, Optional
import re
from pathlib import Path


def parse_text(content: str, filename: str = "") -> str:
    """Parse plain text file"""
    return content.strip()


def parse_markdown(content: str, filename: str = "") -> str:
    """Parse markdown file"""
    # Remove markdown formatting for cleaner text
    text = content

    # Remove code blocks
    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'`[^`]+`', '', text)

    # Remove links but keep text [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    # Remove images
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', text)

    # Remove headers formatting but keep text
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    return text.strip()


def parse_html(content: str, filename: str = "") -> str:
    """Parse HTML file"""
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text.strip()

    except ImportError:
        print("⚠️ BeautifulSoup not installed, falling back to basic HTML parsing")
        # Basic fallback: remove all HTML tags
        text = re.sub(r'<[^>]+>', '', content)
        return text.strip()


def parse_pdf(content: bytes, filename: str = "") -> str:
    """Parse PDF file"""
    try:
        from PyPDF2 import PdfReader
        from io import BytesIO

        pdf_file = BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
            except Exception as e:
                print(f"⚠️ Error extracting page {page_num}: {e}")
                continue

        return "\n\n".join(text_parts)

    except ImportError:
        raise Exception("PyPDF2 not installed. Please install with: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Error parsing PDF: {e}")


def parse_docx(content: bytes, filename: str = "") -> str:
    """Parse DOCX file"""
    try:
        from docx import Document
        from io import BytesIO

        docx_file = BytesIO(content)
        doc = Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    except ImportError:
        raise Exception("python-docx not installed. Please install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {e}")


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks

    Args:
        text: Input text to chunk
        chunk_size: Target size of each chunk in characters
        overlap: Number of characters to overlap between chunks

    Returns:
        List of chunk dictionaries with text and metadata
    """
    if not text.strip():
        return []

    # Split into sentences (basic sentence splitting)
    sentences = re.split(r'(?<=[.!?])\s+', text)

    chunks = []
    current_chunk = ""
    current_size = 0
    chunk_index = 0

    for sentence in sentences:
        sentence_len = len(sentence)

        # If adding this sentence exceeds chunk_size, save current chunk
        if current_size + sentence_len > chunk_size and current_chunk:
            chunks.append({
                "chunk_text": current_chunk.strip(),
                "chunk_index": chunk_index,
                "char_count": len(current_chunk)
            })

            chunk_index += 1

            # Start new chunk with overlap
            if overlap > 0 and len(current_chunk) > overlap:
                # Take last `overlap` characters for overlap
                current_chunk = current_chunk[-overlap:] + " " + sentence
                current_size = len(current_chunk)
            else:
                current_chunk = sentence
                current_size = sentence_len
        else:
            # Add sentence to current chunk
            if current_chunk:
                current_chunk += " " + sentence
            else:
                current_chunk = sentence
            current_size = len(current_chunk)

    # Add final chunk if exists
    if current_chunk.strip():
        chunks.append({
            "chunk_text": current_chunk.strip(),
            "chunk_index": chunk_index,
            "char_count": len(current_chunk)
        })

    return chunks


class DocumentParser:
    """Main document parser that routes to appropriate parser based on file type"""

    SUPPORTED_TYPES = {
        'txt': parse_text,
        'md': parse_markdown,
        'html': parse_html,
        'htm': parse_html,
        'pdf': parse_pdf,
        'docx': parse_docx,
    }

    @classmethod
    def parse(cls, content: Any, file_type: str, filename: str = "") -> str:
        """
        Parse document content based on file type

        Args:
            content: File content (str for text files, bytes for binary)
            file_type: File extension (without dot)
            filename: Original filename

        Returns:
            Extracted text content
        """
        file_type = file_type.lower().replace('.', '')

        if file_type not in cls.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {file_type}. Supported: {list(cls.SUPPORTED_TYPES.keys())}")

        parser_func = cls.SUPPORTED_TYPES[file_type]

        try:
            return parser_func(content, filename)
        except Exception as e:
            raise Exception(f"Error parsing {file_type} file '{filename}': {e}")

    @classmethod
    def parse_and_chunk(
        cls,
        content: Any,
        file_type: str,
        filename: str = "",
        chunk_size: int = 500,
        overlap: int = 50
    ) -> List[Dict[str, Any]]:
        """
        Parse document and split into chunks in one step

        Args:
            content: File content
            file_type: File extension
            filename: Original filename
            chunk_size: Target chunk size
            overlap: Overlap between chunks

        Returns:
            List of chunks with metadata
        """
        # Parse document
        text = cls.parse(content, file_type, filename)

        # Split into chunks
        chunks = chunk_text(text, chunk_size, overlap)

        # Add filename to each chunk
        for chunk in chunks:
            chunk['filename'] = filename

        return chunks
